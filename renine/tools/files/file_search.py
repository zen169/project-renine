"""File search tool for Renine.

Searches for files in the allowed directories by name, extension, or content
(including within PDFs, DOCXs, and text files).
"""
from __future__ import annotations

import os
from pathlib import Path
from typing import Any

import docx
import fitz

from renine.core.config import get_security_config
from renine.core.logging_config import get_logger
from renine.security.input_validator import validate_path
from renine.tools.permissions import PermissionLevel
from renine.tools.registry import BaseTool, ToolResult, register_tool

logger = get_logger(__name__)


def _search_content_in_file(path: Path, query: str) -> bool:
    """Check if the given query text exists inside a file.

    Handles PDFs, DOCXs, and text files. Skips errors.

    Args:
        path: Absolute Path to the file.
        query: Query string to search for (case-insensitive).

    Returns:
        True if the query is found, False otherwise.
    """
    try:
        suffix = path.suffix.lower()
        if suffix == ".pdf":
            with fitz.open(str(path)) as doc:
                for page in doc:
                    if query.lower() in page.get_text().lower():
                        return True
        elif suffix == ".docx":
            doc = docx.Document(path)
            for p in doc.paragraphs:
                if query.lower() in p.text.lower():
                    return True
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if query.lower() in cell.text.lower():
                            return True
        else:
            with open(path, "r", encoding="utf-8", errors="ignore") as f:
                for line in f:
                    if query.lower() in line.lower():
                        return True
    except Exception:
        pass
    return False


def _perform_search(
    base_path: Path,
    query: str,
    search_type: str,
    file_type: str | None = None,
) -> list[dict[str, Any]]:
    """Traverse directory and perform the search based on type and query.

    Args:
        base_path: Base directory Path to search from.
        query: String query.
        search_type: 'name', 'type', or 'content'.
        file_type: Optional extension filter (e.g. '.txt').

    Returns:
        List of matching file info dictionaries.
    """
    results: list[dict[str, Any]] = []
    max_size = get_security_config().get("security", {}).get("filesystem", {}).get("max_read_size_bytes", 52428800)

    for root, _, files in os.walk(base_path):
        for name in files:
            path = Path(root) / name
            suffix = path.suffix.lower()

            # Filter by file extension if requested
            if file_type and suffix != file_type.lower():
                continue

            matched = False
            if search_type == "name" and query.lower() in name.lower():
                matched = True
            elif search_type == "type" and query.lower() in suffix:
                matched = True
            elif search_type == "content":
                # Only check files under policy size limits
                try:
                    if path.stat().st_size <= max_size:
                        matched = _search_content_in_file(path, query)
                except OSError:
                    continue

            if matched:
                try:
                    stat = path.stat()
                    results.append({
                        "name": name,
                        "path": str(path),
                        "size": stat.st_size,
                        "suffix": suffix,
                    })
                except OSError:
                    continue

            # Limit results to prevent output bloat
            if len(results) >= 100:
                break
        if len(results) >= 100:
            break

    return results


@register_tool(
    name="file_search",
    description="Search files by name, type, or content inside allowed directories",
    permission_level=PermissionLevel.READ_ONLY,
    requires_confirmation=False,
)
class FileSearchTool(BaseTool):
    """Tool to search files in the local filesystem."""

    def execute(
        self,
        query: str,
        search_type: str = "name",
        base_dir: str | None = None,
        file_type: str | None = None,
        **kwargs: Any,
    ) -> ToolResult:
        """Search files in a directory.

        Args:
            query: The search term (file name, extension, or content substring).
            search_type: Type of search ('name', 'type', 'content').
            base_dir: Directory path to start search from (defaults to allowed path).
            file_type: Optional file extension to filter by (e.g., '.txt', '.pdf').
            **kwargs: Extra arguments.

        Returns:
            ToolResult with a list of matched files.
        """
        try:
            sec_config = get_security_config().get("security", {})
            allowed = sec_config.get("filesystem", {}).get("allowed_paths", [])

            if base_dir:
                base_path = validate_path(base_dir)
            elif allowed:
                base_path = validate_path(allowed[0])
            else:
                base_path = Path(".").resolve()

            if not base_path.is_dir():
                return ToolResult(
                    success=False,
                    error=f"Base path is not a directory: {base_path}",
                )

            # Extra check: ensure target search folder is allowed
            validate_path(str(base_path))

            results = _perform_search(base_path, query, search_type, file_type)
            return ToolResult(success=True, data={"results": results})

        except Exception as e:
            logger.exception("file_search_failed", query=query, search_type=search_type)
            return ToolResult(success=False, error=str(e))
