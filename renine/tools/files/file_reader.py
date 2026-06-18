"""File reader tool for Renine.

Reads and extracts text from TXT (and other plain text files), DOCX, and PDF
files, performing security and size validation checks.
"""
from __future__ import annotations

from pathlib import Path
from typing import Any

import docx

from renine.core.config import get_security_config
from renine.core.logging_config import get_logger
from renine.security.input_validator import validate_path
from renine.tools.files.pdf_reader import _extract_pdf_text
from renine.tools.permissions import PermissionLevel
from renine.tools.registry import BaseTool, ToolResult, register_tool

logger = get_logger(__name__)


def _extract_docx_text(path: Path) -> str:
    """Extract paragraphs and tables from a DOCX file.

    Args:
        path: Path object to the DOCX file.

    Returns:
        Combined text from document.
    """
    doc = docx.Document(path)
    content: list[str] = []

    for paragraph in doc.paragraphs:
        if paragraph.text:
            content.append(paragraph.text)

    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text for cell in row.cells if cell.text]
            if row_text:
                content.append(" | ".join(row_text))

    return "\n".join(content)


def _read_file_content(path_str: str) -> dict[str, Any]:
    """Helper to validate file and read content based on its extension.

    Args:
        path_str: Path to the target file.

    Returns:
        Dictionary containing content and metadata.
    """
    path = validate_path(path_str)

    if not path.exists():
        raise FileNotFoundError(f"File not found: {path}")

    # Check file size against policy
    sec_config = get_security_config().get("security", {})
    fs_config = sec_config.get("filesystem", {})
    max_size = fs_config.get("max_read_size_bytes", 52428800)

    if path.stat().st_size > max_size:
        raise ValueError(
            f"File size exceeds policy limit of {max_size} bytes."
        )

    suffix = path.suffix.lower()

    if suffix == ".pdf":
        return _extract_pdf_text(path_str)

    if suffix == ".docx":
        text = _extract_docx_text(path)
        return {"text": text, "file_name": path.name, "type": "docx"}

    # Default to text reading
    logger.info("reading_text_file", path=str(path))
    with open(path, "r", encoding="utf-8", errors="replace") as f:
        text = f.read()
    return {"text": text, "file_name": path.name, "type": "text"}


@register_tool(
    name="file_reader",
    description="Read and extract text from TXT, DOCX, or PDF files",
    permission_level=PermissionLevel.READ_ONLY,
    requires_confirmation=False,
)
class FileReaderTool(BaseTool):
    """Tool to read text content from files."""

    def execute(self, file_path: str, **kwargs: Any) -> ToolResult:
        """Extract text from a file.

        Args:
            file_path: Absolute or relative path to the file.
            **kwargs: Extra arguments.

        Returns:
            ToolResult containing the file content.
        """
        if not file_path:
            return ToolResult(success=False, error="File path cannot be empty.")

        try:
            res = _read_file_content(file_path)
            return ToolResult(success=True, data=res)
        except Exception as e:
            logger.exception("file_reading_failed", path=file_path)
            return ToolResult(success=False, error=str(e))
